# -*- coding: utf-8 -*-
"""
Converte GUIA_NEGOCIO_COMPLETO.md para .docx formatado profissionalmente.
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

MD_PATH  = r"C:\Users\Usuario\Documents\ProjetoConcretiza\GUIA_NEGOCIO_COMPLETO.md"
DOC_PATH = r"C:\Users\Usuario\Documents\ProjetoConcretiza\GUIA_NEGOCIO_COMPLETO.docx"

# ── Cores ──────────────────────────────────────────────────────────────────────
C_DARK   = RGBColor(0x0a, 0x0e, 0x27)   # sidebar escura (#0a0e27)
C_ORANGE = RGBColor(0xf9, 0x73, 0x16)   # accent orange (#f97316)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_LIGHT  = RGBColor(0xF5, 0xF5, 0xF5)   # fundo linhas alternadas da tabela
C_GRAY   = RGBColor(0x55, 0x55, 0x55)   # texto secundário

# ── Helpers XML ───────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Define cor de fundo de uma célula."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_table_border(table):
    """Adiciona bordas finas a todas as células."""
    tbl    = table._tbl
    tblPr  = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CCCCCC")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def add_horizontal_rule(doc):
    """Linha horizontal decorativa."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "F97316")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)
    return p


def apply_inline(run_parent, text: str, base_size: int, base_color=None, bold=False):
    """
    Adiciona runs com suporte a **bold**, `code` e links [text](url).
    run_parent pode ser parágrafo ou célula.
    """
    # tokenize: **bold**, `code`, [text](url)
    pattern = r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"
    parts   = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            inner = part[2:-2]
            r = run_parent.add_run(inner)
            r.bold = True
            r.font.size = Pt(base_size)
            if base_color:
                r.font.color.rgb = base_color
        elif part.startswith("`") and part.endswith("`"):
            inner = part[1:-1]
            r = run_parent.add_run(inner)
            r.font.name = "Courier New"
            r.font.size = Pt(base_size - 1)
            r.font.color.rgb = C_ORANGE
        elif re.match(r"\[[^\]]+\]\([^)]+\)", part):
            m     = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            label = m.group(1) if m else part
            r = run_parent.add_run(label)
            r.font.size = Pt(base_size)
            r.font.color.rgb = C_ORANGE
            r.underline = True
        else:
            r = run_parent.add_run(part)
            r.font.size = Pt(base_size)
            if base_color:
                r.font.color.rgb = base_color
            if bold:
                r.bold = True


# ── Estilo do documento ───────────────────────────────────────────────────────

def setup_document() -> Document:
    doc = Document()

    # Margens
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # Fonte padrão
    style         = doc.styles["Normal"]
    font          = style.font
    font.name     = "Calibri"
    font.size     = Pt(10.5)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    return doc


# ── Renderizadores de bloco ───────────────────────────────────────────────────

def render_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    # Fundo escuro simulado com borda esquerda grossa
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "thick")
    left.set(qn("w:sz"),    "24")
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), "F97316")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    r.bold       = True
    r.font.size  = Pt(20)
    r.font.color.rgb = C_DARK
    r.font.name  = "Calibri"
    return p


def render_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold       = True
    r.font.size  = Pt(15)
    r.font.color.rgb = C_ORANGE
    r.font.name  = "Calibri"
    # Linha abaixo
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "F97316")
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def render_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold       = True
    r.font.size  = Pt(12)
    r.font.color.rgb = C_DARK
    r.font.name  = "Calibri"
    return p


def render_h4(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold       = True
    r.italic     = True
    r.font.size  = Pt(11)
    r.font.color.rgb = C_GRAY
    return p


def render_paragraph(doc, text):
    if not text.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    apply_inline(p, text, 10.5)
    return p


def render_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.6)
    p.paragraph_format.space_after  = Pt(3)
    apply_inline(p, text, 10.5)
    return p


def render_numbered(doc, text, num):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    apply_inline(p, text, 10.5)
    return p


def render_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.right_indent = Cm(0.8)
        # fundo cinza claro via shading do parágrafo
        pPr  = p._p.get_or_add_pPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F0F0F0")
        pPr.append(shd)
        r = p.add_run(line if line else " ")
        r.font.name  = "Courier New"
        r.font.size  = Pt(9)
        r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)


def render_blockquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_after  = Pt(6)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "F97316")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    r.italic     = True
    r.font.size  = Pt(10.5)
    r.font.color.rgb = C_GRAY
    return p


def render_table(doc, header_row, sep_row, data_rows):
    """Renderiza tabela markdown com cabeçalho colorido."""
    # Parse colunas
    def parse_cols(row):
        parts = [c.strip() for c in row.strip().strip("|").split("|")]
        return parts

    headers = parse_cols(header_row)
    n_cols  = len(headers)
    n_rows  = len(data_rows)

    if n_cols == 0:
        return

    table = doc.add_table(rows=1 + n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style     = "Table Grid"
    set_table_border(table)

    # Cabeçalho
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], "0A0E27")
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_inline(p, h, 10, base_color=C_WHITE, bold=True)

    # Dados
    for row_idx, row_str in enumerate(data_rows):
        cols      = parse_cols(row_str)
        row_cells = table.rows[row_idx + 1].cells
        # Zebra
        bg = "F5F5F5" if row_idx % 2 == 0 else "FFFFFF"
        for ci in range(n_cols):
            cell_text = cols[ci] if ci < len(cols) else ""
            set_cell_bg(row_cells[ci], bg)
            row_cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = row_cells[ci].paragraphs[0]
            apply_inline(p, cell_text, 10)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Página de rosto ───────────────────────────────────────────────────────────

def add_cover(doc):
    # Título
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.space_after  = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONCRETIZA")
    r.bold = True
    r.font.size  = Pt(36)
    r.font.color.rgb = C_DARK
    r.font.name  = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run("Guia Completo de Negócio")
    r2.bold = True
    r2.font.size  = Pt(22)
    r2.font.color.rgb = C_ORANGE
    r2.font.name  = "Calibri"

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(60)
    r3 = p3.add_run("Do zero ao dia que você vender sua empresa")
    r3.italic = True
    r3.font.size  = Pt(13)
    r3.font.color.rgb = C_GRAY

    add_horizontal_rule(doc)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Março de 2026")
    r4.font.size  = Pt(11)
    r4.font.color.rgb = C_GRAY

    doc.add_page_break()


# ── Parser principal ──────────────────────────────────────────────────────────

def parse_and_render(doc, md_text):
    lines = md_text.splitlines()
    i     = 0
    n     = len(lines)

    while i < n:
        line = lines[i]

        # Página de seção (H1 com número)
        if line.startswith("# ") and re.match(r"^# \d+\.", line):
            text = line[2:].strip()
            doc.add_page_break()
            render_h1(doc, text)
            i += 1
            continue

        # H1 genérico (título, índice, etc.)
        if line.startswith("# "):
            text = line[2:].strip()
            render_h1(doc, text)
            i += 1
            continue

        # H2
        if line.startswith("## "):
            text = line[3:].strip()
            render_h2(doc, text)
            i += 1
            continue

        # H3
        if line.startswith("### "):
            text = line[4:].strip()
            render_h3(doc, text)
            i += 1
            continue

        # H4
        if line.startswith("#### "):
            text = line[5:].strip()
            render_h4(doc, text)
            i += 1
            continue

        # Bloco de código
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            render_code_block(doc, code_lines)
            i += 1
            continue

        # Tabela markdown
        if "|" in line and i + 1 < n and re.match(r"^\s*\|?[-|: ]+\|", lines[i + 1]):
            header_row = line
            sep_row    = lines[i + 1]
            data_rows  = []
            j = i + 2
            while j < n and "|" in lines[j]:
                data_rows.append(lines[j])
                j += 1
            render_table(doc, header_row, sep_row, data_rows)
            i = j
            continue

        # Blockquote
        if line.startswith("> "):
            text = line[2:].strip()
            render_blockquote(doc, text)
            i += 1
            continue

        # Linha horizontal
        if re.match(r"^---+\s*$", line):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Lista com checkbox (- [ ] / - [x])
        if re.match(r"^\s*- \[[ x]\] ", line):
            level  = (len(line) - len(line.lstrip())) // 2
            done   = "[x]" in line[:8]
            text   = re.sub(r"^[\s]*- \[[ x]\] ", "", line)
            prefix = "✅ " if done else "☐ "
            render_bullet(doc, prefix + text, level)
            i += 1
            continue

        # Lista unordered (- ou *)
        if re.match(r"^\s*[-*] ", line):
            level = (len(line) - len(line.lstrip())) // 2
            text  = re.sub(r"^[\s]*[-*] ", "", line)
            render_bullet(doc, text, level)
            i += 1
            continue

        # Lista numbered
        m = re.match(r"^\s*(\d+)\. (.+)", line)
        if m:
            render_numbered(doc, m.group(2), int(m.group(1)))
            i += 1
            continue

        # Linha de separação do índice (ignorar)
        if re.match(r"^\s*$", line):
            i += 1
            continue

        # Parágrafo normal
        render_paragraph(doc, line)
        i += 1


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    print("Lendo markdown...")
    with open(MD_PATH, encoding="utf-8") as f:
        md_text = f.read()

    print("Criando documento Word...")
    doc = setup_document()
    add_cover(doc)
    parse_and_render(doc, md_text)

    print(f"Salvando em {DOC_PATH} ...")
    doc.save(DOC_PATH)
    print("✅ Documento gerado com sucesso!")


if __name__ == "__main__":
    main()
